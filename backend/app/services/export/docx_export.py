"""Render the rich-text editor's HTML into a .docx file.

Supports headings, paragraphs, alignment, lists (nested), block quotes, code
blocks, horizontal rules, tables and inline runs carrying bold / italic /
underline / strike / font family / font size / colour / highlight.
"""
from __future__ import annotations

import base64
import re
from io import BytesIO
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.markdown_tools import markdown_to_html

BLOCK_TAGS = {
    "h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "ul", "ol", "li",
    "blockquote", "pre", "hr", "table", "section", "article",
}
ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
BOLD_WEIGHTS = {"bold", "bolder", "600", "700", "800", "900"}
SIZE_RE = re.compile(r"^([\d.]+)\s*(px|pt|em|rem)?$")
RGB_RE = re.compile(r"rgba?\(([^)]+)\)")


# ------------------------------------------------------------------ style parsing
def _parse_color(value: str) -> RGBColor | None:
    value = (value or "").strip().lower()
    if value.startswith("#"):
        digits = value[1:]
        if len(digits) == 3:
            digits = "".join(ch * 2 for ch in digits)
        if len(digits) == 6:
            try:
                return RGBColor.from_string(digits.upper())
            except ValueError:
                return None
        return None
    match = RGB_RE.match(value)
    if match:
        try:
            parts = [int(float(p)) for p in match.group(1).split(",")[:3]]
            return RGBColor(*parts)
        except (ValueError, TypeError):
            return None
    return None


def _parse_size(value: str) -> Pt | None:
    match = SIZE_RE.match((value or "").strip().lower())
    if not match:
        return None
    number = float(match.group(1))
    unit = match.group(2) or "px"
    if unit == "px":
        number *= 0.75          # CSS pixels -> points
    elif unit in {"em", "rem"}:
        number *= 12.0
    return Pt(max(4.0, min(number, 200.0)))


def _style_attrs(tag: Tag) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for declaration in (tag.get("style") or "").split(";"):
        prop, _, raw = declaration.partition(":")
        prop, raw = prop.strip().lower(), raw.strip()
        if not prop or not raw:
            continue
        if prop == "font-family":
            out["font"] = raw.split(",")[0].strip().strip("'\"")
        elif prop == "font-size":
            size = _parse_size(raw)
            if size:
                out["size"] = size
        elif prop == "color":
            color = _parse_color(raw)
            if color is not None:
                out["color"] = color
        elif prop in {"background-color", "background"}:
            color = _parse_color(raw)
            if color is not None:
                out["shading"] = str(color)
        elif prop == "font-weight":
            out["bold"] = raw in BOLD_WEIGHTS
        elif prop == "font-style":
            out["italic"] = raw == "italic"
        elif prop == "text-decoration" or prop == "text-decoration-line":
            if "underline" in raw:
                out["underline"] = True
            if "line-through" in raw:
                out["strike"] = True
        elif prop == "text-align":
            out["align"] = raw
    return out


TAG_FORMATS: dict[str, dict[str, Any]] = {
    "strong": {"bold": True},
    "b": {"bold": True},
    "em": {"italic": True},
    "i": {"italic": True},
    "u": {"underline": True},
    "ins": {"underline": True},
    "s": {"strike": True},
    "del": {"strike": True},
    "strike": {"strike": True},
    "code": {"font": "Consolas"},
    "kbd": {"font": "Consolas"},
    "mark": {"shading": "FFF3A3"},
    "sub": {"subscript": True},
    "sup": {"superscript": True},
}


def _shade_run(run: Any, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color.replace("#", "").upper())
    run._element.get_or_add_rPr().append(shading)


def _apply(run: Any, fmt: dict[str, Any]) -> None:
    font = run.font
    if fmt.get("bold"):
        run.bold = True
    if fmt.get("italic"):
        run.italic = True
    if fmt.get("underline"):
        run.underline = True
    if fmt.get("strike"):
        font.strike = True
    if fmt.get("subscript"):
        font.subscript = True
    if fmt.get("superscript"):
        font.superscript = True
    if fmt.get("font"):
        font.name = fmt["font"]
    if fmt.get("size"):
        font.size = fmt["size"]
    if fmt.get("color") is not None:
        font.color.rgb = fmt["color"]
    if fmt.get("shading"):
        _shade_run(run, fmt["shading"])


# --------------------------------------------------------------------- rendering
def _add_runs(paragraph: Any, node: Any, inherited: dict[str, Any]) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text.strip() and not text.startswith(" "):
            return
        text = re.sub(r"\s+", " ", text)
        if text:
            _apply(paragraph.add_run(text), inherited)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    if node.name == "img":
        _add_image(paragraph, node)
        return

    fmt = {**inherited, **TAG_FORMATS.get(node.name, {}), **_style_attrs(node)}
    if node.name == "a":
        fmt = {**fmt, "underline": True, "color": _parse_color("#1a56db") or fmt.get("color")}
    for child in node.children:
        _add_runs(paragraph, child, fmt)


def _add_image(paragraph: Any, node: Tag) -> None:
    src = node.get("src") or ""
    if not src.startswith("data:image"):
        return
    try:
        payload = src.split(",", 1)[1]
        stream = BytesIO(base64.b64decode(payload))
        paragraph.add_run().add_picture(stream, width=Inches(5.5))
    except Exception:  # pragma: no cover - malformed inline image
        return


def _paragraph(document: Any, node: Tag, style: str | None, inherited: dict[str, Any]) -> None:
    paragraph = document.add_paragraph(style=style)
    fmt = {**inherited, **_style_attrs(node)}
    align = fmt.pop("align", None)
    if align in ALIGNMENTS:
        paragraph.alignment = ALIGNMENTS[align]
    for child in node.children:
        _add_runs(paragraph, child, fmt)
    if not paragraph.runs and not (node.get_text() or "").strip():
        paragraph.text = ""


def _list_style(ordered: bool, depth: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    return base if depth <= 1 else f"{base} {min(depth, 3)}"


def _render_list(document: Any, node: Tag, ordered: bool, depth: int, inherited: dict[str, Any]) -> None:
    for item in node.find_all("li", recursive=False):
        nested = [child for child in item.find_all(["ul", "ol"], recursive=False)]
        for sub in nested:
            sub.extract()
        _paragraph(document, item, _list_style(ordered, depth), inherited)
        for sub in nested:
            _render_list(document, sub, sub.name == "ol", depth + 1, inherited)


def _render_table(document: Any, node: Tag) -> None:
    rows = node.find_all("tr")
    if not rows:
        return
    width = max(len(row.find_all(["td", "th"])) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row in rows:
        cells = row.find_all(["td", "th"])
        docx_row = table.add_row().cells
        for index in range(width):
            target = docx_row[index]
            target.text = ""
            if index >= len(cells):
                continue
            cell = cells[index]
            paragraph = target.paragraphs[0]
            fmt = {"bold": True} if cell.name == "th" else {}
            fmt.update(_style_attrs(cell))
            align = fmt.pop("align", None)
            if align in ALIGNMENTS:
                paragraph.alignment = ALIGNMENTS[align]
            for child in cell.children:
                _add_runs(paragraph, child, fmt)
    document.add_paragraph()


def _render_block(document: Any, node: Tag, inherited: dict[str, Any]) -> None:
    name = node.name

    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = int(name[1])
        heading = document.add_heading(level=level)
        fmt = {**inherited, **_style_attrs(node)}
        align = fmt.pop("align", None)
        if align in ALIGNMENTS:
            heading.alignment = ALIGNMENTS[align]
        for child in node.children:
            _add_runs(heading, child, fmt)
        return

    if name in {"ul", "ol"}:
        _render_list(document, node, name == "ol", 1, inherited)
        return

    if name == "table":
        _render_table(document, node)
        return

    if name == "blockquote":
        for child in node.children:
            if isinstance(child, Tag) and child.name in BLOCK_TAGS:
                _paragraph(document, child, "Quote", inherited)
            elif isinstance(child, NavigableString) and child.strip():
                document.add_paragraph(str(child).strip(), style="Quote")
        return

    if name == "pre":
        text = node.get_text("\n")
        paragraph = document.add_paragraph()
        _apply(paragraph.add_run(text), {"font": "Consolas", "size": Pt(10)})
        return

    if name == "hr":
        document.add_paragraph("─" * 40)
        return

    if name in {"div", "section", "article"}:
        _render_body(document, node, {**inherited, **_style_attrs(node)})
        return

    _paragraph(document, node, None, inherited)


def _render_body(document: Any, root: Tag, inherited: dict[str, Any]) -> None:
    for child in root.children:
        if isinstance(child, NavigableString):
            if child.strip():
                document.add_paragraph(str(child).strip())
            continue
        if isinstance(child, Tag):
            _render_block(document, child, inherited)


def html_to_docx(html: str, title: str = "") -> bytes:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for section in document.sections:
        section.left_margin = section.right_margin = Inches(1)

    if title:
        document.core_properties.title = title

    soup = BeautifulSoup(html or "", "html.parser")
    body = soup.body or soup
    _render_body(document, body, {})

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def markdown_to_docx(markdown_text: str, title: str = "") -> bytes:
    return html_to_docx(markdown_to_html(markdown_text), title)
